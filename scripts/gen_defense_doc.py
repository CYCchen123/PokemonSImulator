from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

section = doc.sections[0]
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)

style = doc.styles['Normal']
style.font.size = Pt(12)
style.font.name = '宋体'
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

def title(text, size=18):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.bold = True
    r.font.name = '黑体'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.size = Pt(16); r.font.name = '黑体'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.size = Pt(14); r.font.name = '黑体'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

def h3(text):
    p = doc.add_heading(text, level=3)
    for r in p.runs:
        r.font.size = Pt(12); r.font.name = '宋体'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def p(text):
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = Pt(24)
    para.paragraph_format.line_spacing = 1.5
    r = para.add_run(text)
    r.font.size = Pt(12); r.font.name = '宋体'
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def code(text):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(1.5)
    para.paragraph_format.line_spacing = 1.2
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    r = para.add_run(text)
    r.font.size = Pt(9); r.font.name = 'Consolas'

def note(text):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(1)
    r = para.add_run(text)
    r.font.size = Pt(10.5); r.font.name = '楷体'
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

def bullet(text):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(1.5)
    para.paragraph_format.first_line_indent = Cm(-0.5)
    para.paragraph_format.line_spacing = 1.5
    r = para.add_run('• ' + text)
    r.font.size = Pt(12); r.font.name = '宋体'
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ═══════════ 封面 ═══════════
for _ in range(6):
    doc.add_paragraph()
title('宝可梦对战数据分析系统', 22)
title('技术详解说明书', 16)
for _ in range(3):
    doc.add_paragraph()
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p2.add_run('陈奕成  计科2308')
r.font.size = Pt(14); r.font.name = '宋体'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ═══════════ 一 ═══════════
h1('一、数据管道全链路开发')

h2('1.1 背景与设计动机')
p('项目早期采用"本地文件 + 文件监听器"的数据流方案：数据生成器将每场对战写入本地 JSON 文件，后台任务轮询扫描目录，检测新增文件后导入 SQLite。这种方案在小数据量下可以工作，但存在几个问题：')
bullet('扩展性差：文件 I/O 是单机操作，无法利用多节点并行处理能力')
bullet('耦合度高：数据生成、存储、分析全部绑定在同一台机器上')
bullet('数据丢失风险：文件系统不如消息队列可靠')
bullet('无法回放：一旦文件被删除，历史数据无法重新消费')
p('我负责将数据消费端改造为"消息队列 + 流处理"的分布式架构。核心思路是用 Kafka 做数据缓冲和持久化中间件，用 Spark Streaming 做流式消费者，将数据生产和数据消费完全解耦。')

h2('1.2 架构设计')
p('数据由组内其他同学编写的数据生成器产生，以 JSON 格式发送到 Kafka。我负责的是接收端全链路：从 Kafka 消息队列消费、经 Spark Streaming 流处理入库、到 API 查询服务和前端展示。')
code('┌──────────────────────────────────────────────────────────────────┐')
code('│              数据管道全链路（我负责的部分）                        │')
code('│                                                                   │')
code('│  数据生成器(组员)     Kafka(myz:9092)     Spark Streaming         │')
code('│       │        ──→   battle.logs   ──→  analytics_events.py       │')
code('│       │                    │    │               │                 │')
code('│  JSON 事件            持久化  3分区        microbatch消费          │')
code('│  10种事件类型                                      │               │')
code('│                                                    ▼               │')
code('│  浏览器 ←── FastAPI:8000 ←── SQLite ←─────────────┘              │')
code('│  (Windows)    (cyc)          (cyc: ~/data/output.db)              │')
code('└──────────────────────────────────────────────────────────────────┘')

h2('1.3 输入数据格式')
p('数据生成器产生十种 JSON 事件类型，通过 Kafka 发送到 battle.logs Topic。核心事件类型：')
bullet('battle_init（2 条/场）：携带双方完整队伍数据——每只精灵的 speciesID、4 个招式 ID、道具 ID、特性 ID。包含 side_a 和 side_b 两个精灵对象数组')
bullet('turn_damage（2N 条/场）：每回合伤害，携带 target_species、damage、fainted 标志')
bullet('turn_faint：精灵濒死事件')
bullet('battle_result（2 条/场）：胜负结果，携带 winner（"a"/"b"）、result（win/loss）、双方剩余精灵数')
bullet('session_start、turn_executed、turn_switch、turn_heal、turn_ability、page_view 等')

p('这些事件以 JSON 字符串格式写入 Kafka，每个消息一个事件。数据生成器通过 --kafka 参数连接到 Kafka Broker：')
code('python gen_battle_stream.py --interval 2 --kafka 100.107.105.99:9092')

h2('1.4 消息队列：Kafka Topic 设计')
p('Kafka 部署在 myz 节点（100.107.105.99:9092），Topic 名为 battle.logs。关键配置：')
bullet('分区数：3。允许最多 3 个消费者并行读取，匹配 Spark 集群的 Executor 数量')
bullet('副本因子：1。因为集群规模小（3 节点），且数据允许丢失（模拟数据可重新生成），未开启多副本')
bullet('消息 Key：使用 session_id 作为 Key，保证同一会话的事件按顺序进入同一分区')
bullet('消息 Value：JSON 字符串，UTF-8 编码，每条消息一个事件对象')
bullet('保留策略：默认 7 天（604800000ms）。遇到数据积压问题时，临时设为 1 秒触发立即清理')

h2('1.5 为什么选 Kafka')
bullet('持久化：消息写入磁盘，支持数据回放。即使消费者宕机重启，也能从上次 offset 继续消费')
bullet('解耦：生产者（gen_battle_stream）和消费者（Spark）互不知道对方的存在，各自独立扩缩容')
bullet('吞吐量：Kafka 的零拷贝技术和顺序 I/O 保证了高吞吐，单 broker 可达数百 MB/s')
bullet('生态集成：Spark Structured Streaming 原生支持 Kafka Source，一行代码即可集成')

h2('1.6 对比原方案的优势')
bullet('原方案：数据生成器 → JSON 文件 → watcher 轮询 → battle_importer → SQLite')
bullet('新方案：数据生成器 → Kafka → Spark Streaming → SQLite')
p('原方案的瓶颈在 watcher 轮询——每 10 秒扫描全部文件并对比 mtime，数据量大时效率低下。Kafka 的 push 模型消除了轮询开销，Spark 的 microbatch 模式比文件扫描更高效。此外，新方案的每个环节可以独立部署在不同机器上，实现了真正的分布式消费。')

doc.add_page_break()

# ═══════════ 二 ═══════════
h1('二、WebSocket 实时推送服务')

h2('2.1 背景与设计动机')
p('前端需要实时获取最新的排行榜数据。最初采用的是前端定时轮询——每 10 秒发一次 HTTP GET 请求拉取全量数据。这种方式的问题是：')
bullet('无谓开销：即使数据没有变化（没有新对战产生），前端也会发请求')
bullet('延迟不可控：如果数据在两次轮询之间到达，最多有 10 秒延迟')
bullet('服务器负载：多个前端同时轮询产生不必要的并发压力')
p('改进方案是用 WebSocket 实现服务器主动推送：后端检测到数据更新后，立即推送给所有连接的前端。前端只需要保持一个 WebSocket 连接，等待消息即可。')

h2('2.2 后端推送架构')
p('WebSocket 推送服务分为两部分：广播函数和定时广播任务。')

h3('2.2.1 广播函数 broadcast()')
p('broadcast() 是核心推送函数，位于 standalone_server.py 第 613 行。它维护一个全局的 WebSocket 客户端集合 ws_clients，遍历所有连接，发送 JSON 格式的消息：')
code('async def broadcast(msg_type: str, data):')
code('    dead = set()')
code('    for ws in ws_clients:')
code('    try:')
code('        await ws.send_json({"type": msg_type, "data": data})')
code('    except:')
code('        dead.add(ws)                    # 标记断开的连接')
code('    ws_clients.difference_update(dead)   # 清理断开的连接')
p('消息格式为 {"type": "stats_updated", "data": {...}}，type 字段用于前端路由（一条 WebSocket 连接可推送多种消息类型），data 字段携带实际数据。断开的 WebSocket 连接被自动检测并移除。')

h3('2.2.2 集群模式定时广播 _cluster_broadcast_loop()')
p('集群模式下，没有本地文件监听器来触发数据更新推送。因此需要一个独立的定时任务定期查询数据库并广播：')
code('async def _cluster_broadcast_loop():')
code('    logger.info("[cluster] Periodic broadcast started (10s)")')
code('    while True:')
code('        await asyncio.sleep(10)')
code('        try:')
code('            await broadcast("stats_updated", {')
code('                "species_usage":  get_meta_species(),    # 精灵使用率排行')
code('                "move_usage":    get_meta_moves(),       # 招式使用排行')
code('                "item_usage":    get_meta_items(),       # 道具使用排行')
code('                "ability_usage": get_meta_abilities(),   # 特性使用排行')
code('                "type_distribution": get_type_distribution(),')
code('                "summary": get_analysis_summary().get("data", {}),')
code('            })')
code('        except Exception as e:')
code('            logger.error(f"[cluster] Broadcast error: {e}")')

h3('2.2.3 本地模式的文件监听器')
p('本地模式下，推送由文件监听器触发。监听器每 10 秒扫描 battle_logs 目录，检测 JSON 对战文件或 JSONL 事件文件的 mtime 变化。发现新数据后，先调用 battle_importer 或 events_to_db 导入数据库，然后调用 broadcast() 推送：')
code('async def _live_watcher_loop():')
code('    while True:')
code('        await asyncio.sleep(WATCH_INTERVAL)')
code('        latest = _get_latest_battle_mtime(battle_dir)')
code('        if latest > _last_battle_mtime:')
code('            await _run_import_and_broadcast()  # 导入 + 广播')
p('双模式对比：本地模式 = 文件变化触发 + 广播，集群模式 = 定时器触发 + 广播。两种模式共享同一个 broadcast() 函数和同一种消息格式，前端无需区分。')

h2('2.3 前端接收端')
p('前端在 StatsDashboard.vue 的 onMounted 阶段启动 WebSocket 连接：')
code('onMounted(async () => {')
code('    await loadAllData()              # 先加载初始数据')
code('    if (liveMode.value) startLiveMode()  # 连接 WebSocket')
code('})')
p('startLiveMode() 连接 WebSocket 后，注册 stats_updated 消息的回调：')
code('async function startLiveMode() {')
code('    await connect(name)')
code('    liveConnected.value = true')
code('    _unsubStats = on(\'stats_updated\', () => {')
code('        loadStreaming()              # 拉取最新全量数据')
code('    })')
code('}')
p('loadStreaming() 调用 API 端点 /stats/deep/live，返回包含所有维度数据的完整 package，然后更新 speciesRows、moveRows 等响应式变量。Vue 的响应式系统自动触发 DataTable 组件重新渲染。')

h2('2.4 WebSocket 连接路径')
p('本地模式下，API 运行在 localhost:8000，WebSocket URL 为 ws://localhost:8000/ws，直连即可。集群模式下，API 运行在远程 VM（192.168.209.137:8000），前端通过 Vite 开发服务器的代理功能转发 WebSocket 连接：')
code('// vite.config.js')
code('const API_HOST = process.env.VITE_API_HOST || \'http://192.168.209.137:8000\'')
code('proxy: {')
code('    \'/api\': { target: API_HOST, changeOrigin: true },')
code('    \'/ws\':  { target: API_HOST.replace(\'http\', \'ws\'), ws: true, changeOrigin: true }')
code('}')
p('前端连接到 ws://localhost:5173/ws，Vite 服务器将连接转发到 ws://192.168.209.137:8000/ws。对前端来说，API 地址永远是 localhost，无需感知后端实际部署位置。')

h2('2.5 WebSocket 客户端实现（wsClient.js）')
p('前端 WebSocket 客户端是一个单例模块。关键设计：')
bullet('自动重连：连接断开后自动重试')
bullet('消息路由：通过 type 字段分发消息给不同的回调函数（on/off 订阅模式）')
bullet('请求-响应模式：支持带 msgId 的请求，等待匹配的响应消息，实现类似 RPC 的调用')
bullet('连接复用：整个应用共享一个 WebSocket 连接，避免多连接开销')

doc.add_page_break()

# ═══════════ 三 ═══════════
h1('三、Spark Structured Streaming 消费者程序')

h2('3.1 背景与设计动机')
p('Kafka 里的事件是 JSON 字符串，需要通过流处理引擎消费并写入分析数据库。选择 Spark Structured Streaming 而非手动消费的原因是：')
bullet('并行消费：3 个 Kafka 分区可由 3 个 Spark Executor 并行消费，提升吞吐')
bullet('容错性：Spark 的 checkpoint 机制保证 exactly-once 语义——每条消息恰好被处理一次')
bullet('DataFrame API：声明式编程，比手写 Kafka Consumer 循环更简洁')
bullet('分布式能力：消费任务可调度到集群中任意节点执行')

h2('3.2 整体流程')
code('Kafka Topic (battle.logs, 3 分区)')
code('    │')
code('    │ SparkSession.readStream (microbatch, 每触发间隔一批)')
code('    ▼')
code('DataFrame (每行一个 Kafka 消息: key + value + topic + partition + offset)')
code('    │')
code('    │ selectExpr("CAST(value AS STRING) as json_str")')
code('    ▼')
code('DataFrame (每行一个 JSON 字符串)')
code('    │')
code('    │ foreachBatch()')
code('    ▼')
code('process_batch(df, epoch_id)')
code('    │')
code('    ├── _parse_rows(df):  将每行 JSON 字符串 → Python dict')
code('    │')
code('    ├── 读取 dict["event"] 字段，按 EVENT_WRITERS 字典路由')
code('    │')
code('    ├── "battle_init"    → write_battle_init_teams()')
code('    │     └── INSERT 精灵数据 (species_id, ability_id, item_id, move_ids)')
code('    │')
code('    ├── "battle_result"  → write_battle_results()')
code('    │     └── INSERT battles + UPDATE 败方 fainted=1')
code('    │')
code('    ├── "session_start"  → write_users()')
code('    └── "team_save"      → write_user_teams()')

h2('3.3 核心设计决策：为什么不用 Spark SQL from_json()')
p('Spark SQL 提供了 from_json() 函数，可以用 StructType 定义 Schema 来解析 JSON。但本项目的事件有十种类型，每种事件的 data 字段结构完全不同：')
bullet('battle_init 的 data 包含 side_a 和 side_b 两个数组，每个元素是包含 speciesID、moves、item、ability 等字段的嵌套对象')
bullet('turn_damage 的 data 包含 target_species（数字）、damage（数字）、fainted（布尔）')
bullet('battle_result 的 data 包含 result（字符串）、winner（字符串）、turns（数字）')
p('这三种事件的 data 字段结构互不兼容，无法用一个 StructType 统一描述。如果用 from_json，需要定义多个 Schema 并根据 event 字段动态选择——但这会让代码变得复杂。相比之下，直接用 Python 的 json.loads 逐行解析是最简洁的方案：')
code('def _parse_rows(df):')
code('    import json')
code('    result = []')
code('    for r in df.collect():')
code('        try:')
code('            result.append(json.loads(r.json_str))')
code('        except:')
code('            pass')
code('    return result')
p('df.collect() 将当前 microbatch 的所有行收集到 Driver 端。由于每批数据量不大（几十到几百条），不会造成 Driver OOM。每条消息的 JSON 字符串在 json_str 列中，调用 json.loads 即可得到 Python dict。')

h2('3.4 事件路由机制')
p('解析后的每条事件是一个 Python dict，顶层的 event 字段标识事件类型。我设计了一个路由表 EVENT_WRITERS，它是一个字典，将 (事件类型, 目标表) 映射到处理函数：')
code('EVENT_WRITERS = {')
code('    ("battle_init",     "battle_states"): write_battle_init_teams,')
code('    ("battle_result",   "battles"):       write_battle_results,')
code('    ("session_start",   "users"):         write_users,')
code('    ("team_save",       "user_teams"):    write_user_teams,')
code('}')
p('process_batch 函数遍历这个字典，按事件类型分组后分别处理：')
code('def process_batch(df, epoch_id, pokemon_db, output_db):')
code('    rows = _parse_rows(df)')
code('    if not rows:')
code('        return')
code('    for (event_type, target), writer in EVENT_WRITERS.items():')
code('        subset = [r for r in rows if r.get("event") == event_type]')
code('        if subset:')
code('            db = pokemon_db if target in ("users", "user_teams") else output_db')
code('            n = writer(subset, db)')
code('            print(f"  {event_type} → {target}: {n} rows")')
p('注意：users 和 user_teams 写入 pokemon.db（参考数据库），而 battle_init 和 battle_result 的数据写入 output.db（分析数据库）。这里通过 target 字段区分目标库。')

h2('3.5 核心处理函数一：write_battle_init_teams()')
p('这是数据量最大的处理函数，负责处理 battle_init 事件。每个事件携带双方完整队伍数据。')

h3('3.5.1 自动建表')
p('处理函数的第一件事是 CREATE TABLE IF NOT EXISTS，确保目标表存在。建表 SQL 包含 UNIQUE 约束，防止同一场对战的同一精灵被重复插入：')
code('CREATE TABLE IF NOT EXISTS battle_pokemon_states (')
code('    id INTEGER PRIMARY KEY AUTOINCREMENT,')
code('    battle_id TEXT NOT NULL,')
code('    side_index INTEGER DEFAULT 0,        -- 0=A方, 1=B方')
code('    pokemon_index INTEGER DEFAULT 0,     -- 队伍中的槽位 0-5')
code('    species_id INTEGER DEFAULT 0,        -- 宝可梦全国图鉴编号')
code('    hp_pct REAL DEFAULT 100.0,           -- 生命值百分比')
code('    fainted INTEGER DEFAULT 0,           -- 是否濒死 0/1')
code('    ability_id INTEGER DEFAULT 0,        -- 特性ID')
code('    item_id INTEGER DEFAULT 0,           -- 道具ID')
code('    move_ids TEXT DEFAULT \'[]\',          -- 招式列表 (JSON数组)')
code('    slot INTEGER DEFAULT 0,')
code('    UNIQUE(battle_id, species_id, side_index, pokemon_index)')
code(')')
p('UNIQUE 约束是后期优化加入的。在没有约束之前，Spark 重启时会从 Kafka 头部重新消费，导致同一场对战的精灵被 INSERT 多次，部分精灵的出场数从正常的 700 涨到 4000+。加入 UNIQUE 后，INSERT OR IGNORE 会自动跳过重复行。')

h3('3.5.2 数据提取与 INSERT')
p('遍历每个 battle_init 事件，从 data 字段中提取 side_a 和 side_b 两个队伍数组。side_index 固定为 0（A方）或 1（B方）。每个队伍包含 3-6 只精灵，每只精灵是一个 dict，从中提取 speciesID、ability、item、moves 四个核心字段：')
code('def write_battle_init_teams(df_rows, db_path):')
code('    conn = sqlite3.connect(db_path)')
code('    for row in df_rows:')
code('        d = row.get("data", {})            # 事件的 data 对象')
code('        bid = d.get("battle_id", "")       # 对战编号')
code('        for side_key in ("a", "b"):')
code('            side_index = 0 if side_key == "a" else 1')
code('            team = d.get(f"side_{side_key}", [])')
code('            for slot, mon in enumerate(team):')
code('                if not isinstance(mon, dict):')
code('                    continue                   # 跳过非字典元素')
code('                species_id = mon.get("speciesID", 0)')
code('                if not species_id:')
code('                    continue                   # 跳过空槽位')
code('                ability_id = mon.get("ability", 0)')
code('                item_id = mon.get("item", 0)')
code('                moves = json.dumps(mon.get("moves", []))')
code('                conn.execute(')
code('                    "INSERT OR IGNORE INTO battle_pokemon_states"')
code('                    "(battle_id, turn, side_index, pokemon_index,"')
code('                    " species_id, hp, max_hp, hp_pct, fainted,"')
code('                    " ability_id, item_id, move_ids, slot)"')
code('                    "VALUES (?,0,?,?,?,100,100,100.0,0,?,?,?,?)",')
code('                    (bid, side_index, slot, species_id,')
code('                     ability_id, item_id, moves, slot))')
code('    conn.commit()')
code('    conn.close()')
p('初始值设定：hp=100, max_hp=100, hp_pct=100.0, fainted=0。回合数 turn 初始为 0。这些值在后续收到 battle_result 事件时才会更新。每处理完一批事件后 commit() 一次，保证事务一致性。')

h2('3.6 核心处理函数二：write_battle_results()')
p('battle_result 事件携带对战的最终结果。这个函数做了两件事。')

h3('3.6.1 写入 battles 汇总表')
p('将对战结果写入 battles 表，包括双方剩余精灵数、胜负结果等。这个表用于后续的统计查询：')
code('INSERT INTO battles VALUES (')
code('    battle_id, side, player_id, session_id,')
code('    \'battle_result\', result, winner, turns,')
code('    own_remaining, opp_remaining, timestamp')
code(')')

h3('3.6.2 UPDATE 败方濒死标记 — 解决对位胜率 0% 的关键')
p('这是整个消费者程序中最关键的一步。背景是：battle_init 插入时所有精灵的 fainted=0, hp_pct=100。如果不更新 fainted 标志，对位胜率的 SQL 查询（比较双方存活的精灵数量）就无法区分胜负——两边都是全员存活，胜率永远是 0%。')
p('battle_result 事件中的 winner 字段明确标识了胜方（"a" 或 "b"）。由此反推出败方，将败方所有精灵的 fainted 设为 1，hp_pct 设为 0：')
code('def write_battle_results(df_rows, db_path):')
code('    conn = sqlite3.connect(db_path)')
code('    for row in df_rows:')
code('        d = row.get("data", {})')
code('        bid = d.get("battle_id", "")')
code('        winner = d.get("winner", "")')
code('        if bid and winner in ("a", "b"):')
code('            # 败方 = winner 的对面')
code('            loser_idx = 1 if winner == "a" else 0')
code('            conn.execute(')
code('                "UPDATE battle_pokemon_states"')
code('                " SET fainted=1, hp_pct=0.0"')
code('                " WHERE battle_id=? AND side_index=?",')
code('                (bid, loser_idx))')
code('    conn.commit()')
p('这个设计有四两拨千斤的效果：不改表结构、不改对位胜率查询 SQL、不增加额外的 UPDATE 事件处理。仅仅在收到 battle_result 时补一个 UPDATE，就解决了整个对位胜率模块的数据源问题。')

h2('3.7 为什么不用 turn_damage 增量更新 HP')
p('最初考虑过在收到 turn_damage 事件时，用 damage 值逐步减少对应精灵的 hp_pct。但这样做的复杂度远高于收益：')
bullet('turn_damage 事件中只有 damage 数值和 target_species，没有当前 HP 信息，无法精确计算 hp_pct')
bullet('精灵的 max_hp 取决于个体值、努力值、等级等因素，在 events 格式中不可用')
bullet('对位胜率的计算只依赖 fainted 标志（存活数量比较），不依赖 hp_pct 的精确值')
p('因此选择了更简洁的方案：所有精灵初始 hp_pct=100，败方全部清零。HP 显示已从前端移除，只保留胜率百分比。')

h2('3.8 运行模式')
p('Spark 消费者支持两种运行模式：')
bullet('Batch 模式（--batch）：从 Kafka 头部读取所有消息，一次性处理完退出。用于首次导入历史数据或调试。')
bullet('Streaming 模式（默认）：持续运行，从 latest offset 开始消费新消息。生产环境使用此模式。')
p('启动命令：')
code('/opt/bigdata/spark/bin/spark-submit \\')
code('  --master spark://myz:7077 \\')
code('  --packages org.apache.spark:spark-sql-kafka-0-10_2.12:3.5.0 \\')
code('  ~/analytics_events.py \\')
code('  --broker 100.107.105.99:9092 \\')
code('  --pokemon-db ~/pokemon.db \\')
code('  --output-db ~/data/output.db')

h2('3.9 潜在问题与防范')
bullet('重复消费：Spark 重启后如果 checkpoint 丢失，会从 earliest offset 重新消费。通过删除 Kafka 日志段文件和临时调整 retention.ms 清除积压。后续加入 UNIQUE 约束防止重复 INSERT。')
bullet('Driver OOM：_parse_rows 使用 df.collect() 将所有数据收集到 Driver。当前数据量小（每批数百条），不会 OOM。如果数据量增长，可改为 df.toLocalIterator() 逐行处理。')
bullet('SQLite 并发写入：每个处理函数独立打开连接→写入→关闭，避免长时间持锁。process_batch 按顺序调用各 writer（非并行），避免写冲突。')

doc.add_page_break()

# ═══════════ 保存 ═══════════
output = 'D:/poke_simulator/技术详解说明书.docx'
doc.save(output)
print(f'Done: {output}')
