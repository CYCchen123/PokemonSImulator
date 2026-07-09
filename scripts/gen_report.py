from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

section = doc.sections[0]
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.17)
section.right_margin = Cm(3.17)

style = doc.styles['Normal']
style.font.size = Pt(12)
style.font.name = '宋体'
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.35

def add_center(text, size=22, bold=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.name = '黑体'
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

def add_h1(text):
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.size = Pt(16)
        r.font.name = '黑体'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

def add_h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.size = Pt(14)
        r.font.name = '黑体'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

def add_h3(text):
    p = doc.add_heading(text, level=3)
    for r in p.runs:
        r.font.size = Pt(12)
        r.font.name = '宋体'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_p(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.name = '宋体'
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_item(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.name = '宋体'
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ═══════ 封面 ═══════
for _ in range(6):
    doc.add_paragraph()
add_center('2026年计算机专业生产实习报告', 22)
for _ in range(3):
    doc.add_paragraph()
info = [
    '姓    名：______陈奕成______',
    '学    号：____________________',
    '班    级：______计科2308______',
    '实习时间：2026年6月22日 ~ 2026年7月17日',
    '指导老师：____________________',
    '实习单位：大数据与人工智能综合实训平台',
    '实习项目：基于分布式流处理平台的宝可梦对战数据分析系统',
]
for item in info:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(item)
    r.font.size = Pt(14)
    r.font.name = '宋体'
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ═══════ 目录 ═══════
doc.add_page_break()

# ═══════ 正文 ═══════

# ── 1 ──
add_h1('1  实习目的与意义')
add_h2('1.1 实习目的')
add_p('本次生产实习旨在深入理解和实践大数据分布式处理技术，通过构建"宝可梦对战数据分析系统"这一完整工程项目，系统性地掌握从数据采集、消息队列传输、流式处理到前端可视化的全链路技术体系。')
items = [
    '理解并实践 Apache Kafka 分布式消息队列的部署架构、Topic 设计以及生产者/消费者模式；',
    '掌握 Apache Spark Structured Streaming 的实时流处理编程，实现对对战数据的近实时聚合分析；',
    '综合运用 Vue 3 前端框架与 Python FastAPI 后端框架，构建具有良好用户体验的数据可视化仪表盘；',
    '实践 SQLite 数据库在实时分析场景下的查询优化策略，包括持久化连接池、TTL 缓存等性能调优技术；',
    '掌握 ZooKeeper + Kafka + Spark 多节点 Linux 集群的部署、运维与故障排查能力。',
]
for item in items:
    add_item(item)

add_h2('1.2 实习意义')
add_p('随着大数据技术在互联网行业的广泛应用，实时数据处理与分析已成为企业决策和产品优化的核心能力。宝可梦（Pokémon）作为全球知名的游戏 IP，其对战系统涉及复杂的数值平衡和 meta 演化，天然适合作为数据分析的研究对象。通过本项目的开发，将课堂所学的分布式计算、数据库原理、软件工程、前端开发等理论知识转化为实际工程能力，完整经历"需求分析—系统设计—编码实现—测试部署"的软件生命周期，为今后从事大数据分析、后端开发或全栈工程师岗位奠定扎实的实践基础。')

# ── 2 ──
add_h1('2  实习任务')
add_p('本次实习的核心任务是设计并实现一套完整的宝可梦对战数据实时分析系统，具体包括：')
tasks = [
    '数据分析引擎搭建：设计并实现基于 SQLite 持久化连接 + 内存 TTL 缓存的分析查询引擎，支持精灵使用率排行、招式排行、道具排行、特性排行、属性分布、同队搭配分析、对位胜率查询等 10+ 种统计指标，查询延迟控制在毫秒级别。',
    '实时数据管道构建：搭建 Kafka → Spark Structured Streaming → SQLite 的端到端实时数据管道，实现数据从生成到入库的全链路自动化，支持每秒数条对战事件的吞吐。',
    '前端可视化开发：基于 Vue 3 + Chart.js 实现分析仪表盘，包含带 FLIP 平滑动画的排行榜组件、精灵详情钻取面板、对位胜率 PK 对比界面、同队搭配图表等交互功能。',
    '集群环境部署：在三节点 Ubuntu Server 22.04 虚拟机集群（myz / cyc / lzx）上部署 ZooKeeper 3 节点 Quorum 集群 + Kafka Broker + Spark Master/Worker 分布式计算集群。',
    '双模式兼容设计：通过环境变量 POKEMON_MODE 实现本地单机模式与分布式集群模式的灵活切换，共享同一套前端和 API 路由代码。',
]
for t in tasks:
    add_item(t)

# ── 3 ──
add_h1('3  实习内容')
add_h2('3.1 需求分析')
add_p('系统核心需求围绕宝可梦 Gen9（第九世代）对战数据的实时采集与多维度分析展开：')
reqs = [
    '精灵使用率排行：统计各精灵在所有对战中的出场次数、胜率，按出场数降序排列；',
    '招式/道具/特性排行：分别统计招式使用频次、道具携带率、特性携带率；',
    '实时排行变化动画：前端排名变化时通过 FLIP 动画平滑过渡，提升数据可视化体验；',
    '同队搭配分析：找出经常出现在同一队伍中的精灵组合（自连接 + GROUP BY 聚合）；',
    '对位胜率查询：输入两个精灵，查询它们在历史对战中的对位胜率、平均存活率等指标；',
    '精灵详情钻取：点击排行榜中任意精灵，展开显示该精灵最常用的招式、道具、特性分布；',
    '双模式兼容：本地开发模式（SQLite 直连 + 文件监听）与分布式集群模式（Kafka + Spark）灵活切换；',
    '远程产数控制：通过前端按钮远程控制 VM 集群上的数据模拟器启停。',
]
for r in reqs:
    add_item(r)

add_h2('3.2 概要设计')
add_p('系统采用三层架构：展示层（Vue 3 前端，Windows 物理机）、服务层（Python FastAPI，cyc 节点）、数据层（Kafka + Spark + SQLite，三节点 VM 集群）。')
add_p('展示层：基于 Vue 3 + Vite 构建 SPA 应用，包含 5 个标签页（Overview、Rankings、Battle Data、Types、Scout）。核心组件 DataTable 实现双表架构——filtered 维护数据准确性，_display 独立负责动画缓冲。通过 Vite 开发服务器的代理功能跨网络访问 VM 集群上的 API 服务。')
add_p('服务层：运行于 cyc 节点（192.168.209.137:8000），FastAPI 框架提供 RESTful API + WebSocket 实时推送。核心模块包括 sql_analytics.py（SQL 分析引擎，持久化连接 + TTL 缓存）、battle_importer.py（数据导入器，自动建表 + 增量导入）、standalone_server.py（API 路由 + 后台任务管理）。')
add_p('数据层：myz 节点运行 Kafka Broker（100.107.105.99:9092），battle.logs Topic 3 分区接收对战事件。Spark Structured Streaming 运行 analytics_events.py 消费消息，经 JSON 解析和事件类型路由（battle_init → INSERT 完整队伍数据，battle_result → UPDATE 败方 fainted 标记），写入 cyc 节点的 SQLite 数据库（~/data/output.db）。ZooKeeper 三节点 Quorum 集群提供分布式协调。')

add_h2('3.3 详细设计')
add_p('数据库设计：核心表 battle_pokemon_states 包含 battle_id、turn、side_index、species_id、hp_pct、fainted、ability_id、item_id、move_ids（JSON 数组）等 14 个字段。精灵参考数据（名称、属性、种族值）存储于独立的 pokemon.db 参考库。查询时通过 Python 侧 join 实现跨库关联。')
add_p('API 设计：共 15+ 个 RESTful 端点，覆盖五大功能模块。核心端点包括 GET /stats/deep/summary（统计摘要，含总对战数、总精灵种类数）、GET /stats/deep/meta（精灵使用率排行）、GET /stats/deep/head-to-head?s1=&s2=（对位胜率查询）、GET /stats/deep/team-synergy（同队搭配）、POST /cluster/gen/start|stop（远程产数控制）。WebSocket 端点 /ws 用于服务器主动推送 stats_updated 事件。')
add_p('前端 DataTable 动画设计：双表分离架构是本项目的核心技术创新。filtered（表A）在每次数据到达时立即更新为正确值，确保查询结果始终准确。_display（表B）从旧排名快照出发，通过两步 FLIP 流程实现视觉动画——第一步设置旧顺序让 Vue 记录 DOM 位置（before），第二步设置目标顺序触发 FLIP 计算 transform 并应用 CSS transition（1 秒）。动画结束后 _display 置为 null，前端切回 filtered 渲染。旧 tick 被新数据中断时只退出不碰 _display，新 tick 无缝接管。')

# ── 4 ──
add_h1('4  项目开发环境')
add_h2('4.1 硬件环境')
add_p('开发机：Windows 11 Home China 物理机（16GB RAM, SSD），运行前端 Vite 开发服务器。')
add_p('集群：三台 VMware 虚拟机组成 Ubuntu Server 22.04 集群，通过 Tailscale 虚拟网络（100.x CIDR）互联。')
hw = [
    'myz（192.168.88.129 / Tailscale 100.107.105.99）：Kafka Broker、Spark Master、ZooKeeper',
    'cyc（192.168.209.137 / Tailscale 100.123.146.105）：API Server、Spark Worker、ZooKeeper、数据生成器',
    'lzx（Tailscale 100.109.69.122）：Spark Worker、ZooKeeper Follower',
]
for h in hw:
    add_item(h)

add_h2('4.2 软件环境')
sw = [
    'Python 3.12 + FastAPI 0.139 → 后端 API 服务',
    'Vue 3 + Vite 5 → 前端 SPA 框架',
    'Apache Kafka 3.4.1 → 分布式消息队列',
    'Apache Spark 3.5.0 → 分布式流处理引擎',
    'Apache ZooKeeper 3.8.6 → 分布式协调服务',
    'SQLite 3 → 分析数据库',
    'Chart.js 4.x → 数据可视化图表库',
    'Tailscale → 虚拟组网',
    'VMware Workstation → 虚拟机管理',
]
for s in sw:
    add_item(s)

# ── 5 ──
add_h1('5  项目开发过程')
add_h2('5.1 第一阶段：本地单机版本原型开发')
add_p('项目初期聚焦于快速搭建可用的分析原型。以 Python FastAPI 为后端框架，配合 Vue 3 前端，构建了基础的数据查询与展示功能。后端核心模块 sql_analytics.py 实现了 SQLite 持久化连接与线程本地存储（thread-local），配合 0.5 秒 TTL 缓存机制，将常见查询的延迟控制在 1 毫秒以内。数据产生方面，开发了 gen_battle_stream.py 对战模拟器，从 pokemon.db 参考库中加载真实的精灵、招式、特性、道具数据，每 2 秒生成一场包含完整队伍信息的对战事件流，通过文件监听器（watcher）每 10 秒扫描 battle_logs 目录并自动导入。前端实现了 5 标签页仪表盘，每个标签页包含 Top 10 排行榜、Chart.js 图表组件。')

add_h2('5.2 第二阶段：前端动画优化与交互增强')
add_p('为提升排行变化的视觉体验，对 DataTable 组件进行了重点优化。引入 Vue 3 TransitionGroup FLIP 动画机制，实现排名变化时行的平滑滑动效果。核心创新在于"双表分离"架构——filtered 在数据到达时立即设为正确值，_display 从旧快照通过 requestAnimationFrame → 目标顺序 → CSS transition 1s 的流程实现动画，两张表完全解耦。同时引入幽灵占位（始终在分页末补足 10 行空占位）保证表格高度恒定，避免 TransitionGroup 动画导致页面抖动。还开发了精灵详情钻取面板、同队搭配热力图、对位胜率 PK 对比界面。')

add_h2('5.3 第三阶段：分布式集群部署与管道构建')
add_p('在三台 VM 上部署完整的分布式栈：ZooKeeper 3 节点 Quorum 集群（容忍单节点故障）、Kafka Broker（myz 节点，3 分区 Topic）、Spark 集群（Master myz + Worker cyc/lzx）。数据管道中 gen_battle_stream.py 作为 Kafka Producer 发送 10 种事件类型（battle_init、turn_damage、turn_faint、battle_result 等），Spark Structured Streaming 通过 analytics_events.py 消费并进行事件路由（battle_init → INSERT 完整队伍数据含招式/道具/特性，battle_result → UPDATE 败方 fainted 标记），写入 SQLite 分析数据库。')

add_h2('5.4 第四阶段：双模式兼容与系统联调')
add_p('通过环境变量 POKEMON_MODE=local|cluster 控制双模式切换。本地模式使用 SQLite 直连 + 文件监听，适合开发调试。集群模式启用 Kafka + Spark 管道 + WebSocket 每 10 秒广播 stats_updated 推送。联调中解决了 Kafka advertised.listeners 配置错误、ZooKeeper 单点故障、Kafka 历史数据积压导致重复消费、WebSocket 代理配置错误等关键技术问题。编写了 start.sh（一键启动集群）、clear_data.sh（清理 DB 并显示清除条数）、check_dups.sh（重复数据诊断）等运维脚本，实现集群的便捷管理。')

# ── 6 ──
add_h1('6  项目总结')
add_h2('6.1 技术成果')
results = [
    'SQLite 查询优化引擎：持久化连接 + 线程本地存储 + 0.5s TTL 内存缓存，查询延迟 <1ms；',
    'Vue 3 FLIP 动画排行榜：双表分离 + 幽灵占位 + 状态机控制，1s 平滑过渡无中间态错误；',
    'Kafka → Spark → SQLite 分布式管道：支持每秒多条事件流处理，容量可横向扩展；',
    '多维度分析：10+ 种统计维度的 RESTful API（含对位胜率、同队搭配）；',
    '双模式兼容：一套代码支持 local 开发模式和 cluster 生产模式；',
    '一键运维：start.sh / clear_data.sh / check_dups.sh 等脚本实现集群便捷管理。',
]
for r in results:
    add_item(r)

add_h2('6.2 遇到的问题与解决方案')

add_h3('6.2.1 Kafka 连接超时')
add_p('问题：Kafka advertised.listeners 配置指向不可达主机名，客户端初始连接后被重定向到错误 IP，60 秒超时。')
add_p('解决：将 advertised.listeners 配置为 Tailscale 虚拟网络 IP（100.107.105.99:9092），确保集群内所有节点可达。')

add_h3('6.2.2 ZooKeeper 集群失联')
add_p('问题：三节点 ZK 集群中两台关机，单节点无法形成 Quorum（需 >50% 在线），导致 Kafka 启动失败。')
add_p('解决：建立启动检查清单，每次先确认至少 2 台 ZK 节点在线（zkServer.sh status），再启动 Kafka。')

add_h3('6.2.3 排行榜动画中间态数据错误')
add_p('问题：逐行 splice 动画使用预计算静态 newIdx，但 splice 后数组变化导致后续行目标位置失准，出现 1 出场精灵错误排在 2 出场精灵之间的中间态。')
add_p('解决：放弃逐行 splice，改为 FLIP 两步到位法——_display 旧顺序 → Vue 记录位置 → _display 目标顺序，Vue FLIP 全自动计算 transform 过渡，不存在错误中间态。')

add_h3('6.2.4 Spark 重复消费导致数据膨胀')
add_p('问题：Spark 重启后从 Kafka earliest offset 重新消费全部历史消息（112MB, 165 个日志段），同一场对战被反复插入，导致部分精灵出场数异常偏高（4000+ vs 预期 700）。')
add_p('解决：调整 Kafka topic retention.ms 至 1 秒清理历史消息，并删除 /opt/bigdata/data/kafka/battle.logs* 日志段文件彻底清除积压。')

add_h3('6.2.5 集群模式下前端无数据')
add_p('问题：前端 WebSocket URL 使用 ws://localhost:8000/ws 直连 API 端口，但 API 运行于远程 VM 而非本地，握手超时无法接收 stats_updated 推送。')
add_p('解决：修改 wsClient.js 的 WS_URL 通过 Vite 开发服务器代理转发 WebSocket；在 vite.config.js 中将 API_HOST 默认值设为 VM 地址（192.168.209.137:8000）。')

add_h2('6.3 心得体会')
add_p('通过本次生产实习，我深刻体会到理论知识与工程实践之间的差距。在课堂上学习 Kafka、Spark 等分布式技术时，了解的是它们的概念和 API 用法；但在实际部署中，网络配置（advertised.listeners）、进程协调（ZooKeeper Quorum）、数据一致性（重复消费去重）等工程细节才是决定系统能否稳定运行的关键。')
add_p('在前端开发中，我认识到视觉效果与数据正确性往往是相互冲突的目标。排行榜动画的"回退"bug 本质上是数据流与 UI 流同步的经典问题——通过引入"双表分离"架构，将数据层与动画层解耦，才从根本上消除了竞态条件。这一设计思路可推广到其他需要实时数据展示的场景。')
add_p('此外，本次实习也锻炼了我的跨领域问题排查能力。从 Kafka 连接超时到 Spark 消费积压，每个问题都需要综合运用 Linux 系统管理、网络原理、分布式系统理论来定位和解决。这种端到端的问题解决经验是课堂教学难以提供的。')

# ── 参考文献 ──
add_h1('参考文献')
refs = [
    '[1] Apache Software Foundation. Apache Kafka Documentation[EB/OL]. https://kafka.apache.org/documentation/, 2026.',
    '[2] Apache Software Foundation. Spark Structured Streaming Programming Guide[EB/OL]. https://spark.apache.org/docs/latest/structured-streaming-programming-guide.html, 2026.',
    '[3] Evan You. Vue.js TransitionGroup API[EB/OL]. https://vuejs.org/guide/built-ins/transition-group.html, 2026.',
    '[4] Sebastián Ramírez. FastAPI Documentation[EB/OL]. https://fastapi.tiangolo.com/, 2026.',
    '[5] Apache Software Foundation. ZooKeeper Administrator\'s Guide[EB/OL]. https://zookeeper.apache.org/doc/current/zookeeperAdmin.html, 2026.',
    '[6] Pokémon Showdown. Pokémon Damage Calculator & Simulator[EB/OL]. https://pokemonshowdown.com/, 2026.',
    '[7] Neha Narkhede et al. Kafka: The Definitive Guide[M]. O\'Reilly Media, 2017.',
    '[8] Holden Karau et al. Learning Spark: Lightning-Fast Big Data Analysis[M]. O\'Reilly Media, 2015.',
]
for ref in refs:
    p = doc.add_paragraph()
    r = p.add_run(ref)
    r.font.size = Pt(10.5)
    r.font.name = '宋体'
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

output = 'D:/poke_simulator/2026年计算机专业生产实习报告.docx'
doc.save(output)
print(f'Done: {output}')
